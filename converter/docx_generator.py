"""
DOCX 生成器：从中间表示生成 DOCX 文档
"""
import re
from typing import Dict, Any, List, Tuple, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml

from utils.text_cleanup import cleanup_text, convert_quotes
from latex2mathml.converter import convert as latex_to_mathml


class FontSettings:
    """字体设置"""
    def __init__(self):
        self.chinese_font = "宋体"
        self.english_font = "Times New Roman"
        self.font_size = 12  # 磅值


class ParagraphSettings:
    """段落设置"""
    def __init__(self):
        # 行间距
        self.line_spacing = 1.5  # 单倍、1.5倍、2倍或具体值
        self.line_spacing_type = WD_LINE_SPACING.MULTIPLE

        # 首行缩进
        self.first_line_indent_enabled = True
        self.first_line_indent = 2  # 字符数

        # 标题设置
        self.heading_line_spacing = 1.5
        self.heading_space_before = Pt(12)
        self.heading_space_after = Pt(6)


class DocxGenerator:
    """DOCX 文档生成器"""

    def __init__(self, font_settings: FontSettings, paragraph_settings: ParagraphSettings):
        self.font_settings = font_settings
        self.paragraph_settings = paragraph_settings
        self.document = Document()

    def set_document_properties(self):
        """设置文档默认属性"""
        self.document.styles['Normal'].font.name = self.font_settings.english_font
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                                 self.font_settings.chinese_font)
        self.document.styles['Normal'].font.size = Pt(self.font_settings.font_size)

    def apply_run_font(self, run, text: str):
        """为 run 设置字体，根据文本内容混合中英文"""
        has_chinese = bool(re.search(r'[一-龥]', text))
        if has_chinese:
            run.font.name = self.font_settings.english_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_settings.chinese_font)
        else:
            run.font.name = self.font_settings.english_font

    def create_paragraph_with_format(self, text: str, style: Optional[str] = None,
                                     alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
        text = cleanup_text(text)
        text = convert_quotes(text)
        parts = self._split_inline_formulas(text)

        if style:
            p = self.document.add_paragraph(style=style)
        else:
            p = self.document.add_paragraph()

        p.alignment = alignment
        self._apply_paragraph_format(p)

        for part in parts:
            if part['type'] == 'text':
                if part['text'].strip():
                    run = p.add_run(part['text'])
                    self.apply_run_font(run, part['text'])
            elif part['type'] == 'formula':
                self._add_inline_formula(p, part['formula'])

    def _split_inline_formulas(self, text: str) -> List[Dict[str, str]]:
        parts = []
        pattern = r'(\$[^$]+\$)'
        fragments = re.split(pattern, text)

        for fragment in fragments:
            if fragment.startswith('$') and fragment.endswith('$') and len(fragment) > 2:
                formula = fragment[1:-1]
                parts.append({'type': 'formula', 'formula': formula})
            elif fragment.strip():
                parts.append({'type': 'text', 'text': fragment})

        return parts

    def _add_inline_formula(self, paragraph, formula: str) -> None:
        """添加行内公式"""
        try:
            mathml_xml = latex_to_mathml(formula.strip())
            mathml_elem = etree.fromstring(mathml_xml)

            # 创建公式的run
            run = paragraph.add_run()
            r = run._r

            # 设置字体
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Cambria Math')
            rFonts.set(qn('w:hAnsi'), 'Cambria Math')
            rFonts.set(qn('w:eastAsia'), 'Cambria Math')
            rPr.append(rFonts)
            r.insert(0, rPr)

            # 创建oMath元素
            oMath = OxmlElement('m:oMath')

            # 转换MathML到OMML
            self._mathml_to_omml(mathml_elem, oMath)

            # 添加到run
            r.append(oMath)

        except Exception as e:
            # 如果失败，添加纯文本
            run = paragraph.add_run(f'[{formula}]')
            run.font.name = 'Cambria Math'

    def _mathml_to_omml(self, mathml_elem, oMath):
        """将MathML元素转换为OMML"""
        MATHML_NS = 'http://www.w3.org/1998/Math/MathML'

        for child in mathml_elem:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag in ('mrow', 'math'):
                self._mathml_to_omml(child, oMath)
            elif tag == 'mi':  # 标识符
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mn':  # 数字
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mo':  # 操作符
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mtext':  # 文本
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mspace':  # 空格
                pass
            elif tag == 'msqrt':  # 平方根
                self._convert_msqrt_simple(child, oMath)
            elif tag == 'mroot':  # 根号
                self._convert_mroot_simple(child, oMath)
            elif tag == 'mfrac':  # 分数
                self._convert_mfrac_simple(child, oMath)
            elif tag == 'msub':  # 下标
                self._convert_msub_simple(child, oMath)
            elif tag == 'msup':  # 上标
                self._convert_msup_simple(child, oMath)
            elif tag == 'msubsup':  # 上下标
                self._convert_msubsup_simple(child, oMath)
            elif tag == 'munder':  # 下标/下划线
                self._convert_munder_simple(child, oMath)
            elif tag == 'mover':  # 上标/上划线
                self._convert_mover_simple(child, oMath)
            elif tag == 'munderover':  # 上下组合
                self._convert_munderover_simple(child, oMath)
            elif tag == 'mrow':  # 已经是mrow但上面没处理
                self._mathml_to_omml(child, oMath)

    def _create_omml_run(self, text: str) -> OxmlElement:
        """创建OMML run元素"""
        r = OxmlElement('m:r')
        rPr = OxmlElement('m:rPr')
        rFonts = OxmlElement('m:rFonts')
        rFonts.set(qn('w:ascii'), 'Cambria Math')
        rFonts.set(qn('w:hAnsi'), 'Cambria Math')
        rFonts.set(qn('w:eastAsia'), 'Cambria Math')
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('m:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        return r

    def _convert_msqrt_simple(self, msqrt_elem, oMath):
        """转换平方根"""
        rad = OxmlElement('m:rad')

        # 次数（默认2）
        deg = OxmlElement('m:deg')
        deg.append(self._create_omml_run('2'))
        rad.append(deg)

        # 内容
        e = OxmlElement('m:e')
        for child in msqrt_elem:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'mrow':
                self._mathml_to_omml(child, e)
        rad.append(e)
        oMath.append(rad)

    def _convert_mroot_simple(self, mroot_elem, oMath):
        """转换根号"""
        rad = OxmlElement('m:rad')
        children = list(mroot_elem)

        if len(children) >= 2:
            # 次数
            deg = OxmlElement('m:deg')
            self._mathml_to_omml(children[1], deg)
            rad.append(deg)

            # 内容
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            rad.append(e)
            oMath.append(rad)

    def _convert_mfrac_simple(self, mfrac_elem, oMath):
        """转换分数"""
        f = OxmlElement('m:f')
        children = list(mfrac_elem)

        if len(children) >= 2:
            # 分子
            num = OxmlElement('m:num')
            self._mathml_to_omml(children[0], num)
            f.append(num)

            # 分母
            den = OxmlElement('m:den')
            self._mathml_to_omml(children[1], den)
            f.append(den)

            e = OxmlElement('m:e')
            e.append(f)
            oMath.append(e)

    def _convert_msub_simple(self, msub_elem, oMath):
        """转换下标"""
        children = list(msub_elem)

        if len(children) >= 2:
            sSub = OxmlElement('m:sSub')

            # 基底
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            sSub.append(e)

            # 下标
            sub = OxmlElement('m:sub')
            self._mathml_to_omml(children[1], sub)
            sSub.append(sub)

            oMath.append(sSub)

    def _convert_msup_simple(self, msup_elem, oMath):
        """转换上标"""
        children = list(msup_elem)

        if len(children) >= 2:
            sSup = OxmlElement('m:sSup')

            # 基底
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            sSup.append(e)

            # 上标
            sup = OxmlElement('m:sup')
            self._mathml_to_omml(children[1], sup)
            sSup.append(sup)

            oMath.append(sSup)

    def _convert_msubsup_simple(self, msubsup_elem, oMath):
        """转换上下标"""
        children = list(msubsup_elem)

        if len(children) >= 3:
            sSubSup = OxmlElement('m:sSubSup')

            # 基底
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            sSubSup.append(e)

            # 下标
            sub = OxmlElement('m:sub')
            self._mathml_to_omml(children[1], sub)
            sSubSup.append(sub)

            # 上标
            sup = OxmlElement('m:sup')
            self._mathml_to_omml(children[2], sup)
            sSubSup.append(sup)

            oMath.append(sSubSup)

    def _convert_munder_simple(self, munder_elem, oMath):
        """转换下划线"""
        children = list(munder_elem)

        if len(children) >= 2:
            acc = OxmlElement('m:acc')
            acc.set(qn('m:val'), '_')

            # 基底
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            acc.append(e)

            oMath.append(acc)

    def _convert_mover_simple(self, mover_elem, oMath):
        """转换上划线"""
        children = list(mover_elem)

        if len(children) >= 2:
            acc = OxmlElement('m:acc')
            acc.set(qn('m:val'), '‾')  # 上划线

            # 基底
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            acc.append(e)

            oMath.append(acc)

    def _convert_munderover_simple(self, munderover_elem, oMath):
        """转换上下组合"""
        children = list(munderover_elem)

        if len(children) >= 3:
            acc = OxmlElement('m:acc')
            acc.set(qn('m:val'), '‾')

            # 基底
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            acc.append(e)

            # 下标
            sub = OxmlElement('m:sub')
            self._mathml_to_omml(children[1], sub)
            acc.append(sub)

            # 上标
            sup = OxmlElement('m:sup')
            self._mathml_to_omml(children[2], sup)
            acc.append(sup)

            oMath.append(acc)

    def _apply_paragraph_format(self, paragraph) -> None:
        """应用段落格式（内部方法）"""
        settings = self.paragraph_settings
        # 设置行间距
        if settings.line_spacing_type == WD_LINE_SPACING.MULTIPLE:
            paragraph.paragraph_format.line_spacing = settings.line_spacing
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        else:
            paragraph.paragraph_format.line_spacing_rule = settings.line_spacing_type
            paragraph.paragraph_format.line_spacing = settings.line_spacing

        # 正文段落设置
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if settings.first_line_indent_enabled:
            # 首行缩进：1 字符 ≈ 字号(pt) 宽度
            indent_value = Pt(settings.first_line_indent * self.font_settings.font_size)
            paragraph.paragraph_format.first_line_indent = indent_value

    def add_heading(self, text: str, level: int) -> None:
        text = cleanup_text(text)
        text = convert_quotes(text)

        p = self.document.add_heading(text, level=level)
        for run in p.runs:
            run.font.name = self.font_settings.english_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'),
                                        self.font_settings.chinese_font)

        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing = self.paragraph_settings.heading_line_spacing
        paragraph_format.space_before = self.paragraph_settings.heading_space_before
        paragraph_format.space_after = self.paragraph_settings.heading_space_after

    def add_paragraph(self, text: str) -> None:
        self.create_paragraph_with_format(text)

    def add_block_formula(self, formula: str) -> None:
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mml = self._latex_to_omml(formula)
        run = p.add_run()
        run._r.append(mml)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Cambria Math')
        rFonts.set(qn('w:hAnsi'), 'Cambria Math')
        rFonts.set(qn('w:eastAsia'), 'Cambria Math')
        p.paragraph_format.space_after = Pt(12)

    def _latex_to_omml(self, latex: str) -> OxmlElement:
        """将LaTeX公式转换为OMML"""
        try:
            mathml_xml = latex_to_mathml(latex.strip())
            mathml_elem = etree.fromstring(mathml_xml)
            oMath = OxmlElement('m:oMath')
            self._mathml_to_omml(mathml_elem, oMath)
            oMathPara = OxmlElement('m:oMathPara')
            oMathPara.append(oMath)
            return oMathPara
        except Exception as e:
            return self._create_fallback_formula(latex)

    def _create_fallback_formula(self, latex: str) -> OxmlElement:
        """创建备用公式（当转换失败时）"""
        oMathPara = OxmlElement('m:oMathPara')
        oMath = OxmlElement('m:oMath')
        r = OxmlElement('m:r')
        t = OxmlElement('m:t')
        t.text = latex
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        oMath.append(r)
        oMathPara.append(oMath)
        return oMathPara

    def add_code_block(self, code: str, language: str) -> None:
        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._p.get_or_add_pPr().append(shading)

    def add_list_item(self, text: str) -> None:
        self.create_paragraph_with_format(f"• {text}")

    def add_blockquote(self, text: str) -> None:
        text = cleanup_text(text)
        text = convert_quotes(text)

        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        run.font.italic = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '808080')
        pBdr.append(left)
        pPr.append(pBdr)

    def add_horizontal_rule(self) -> None:
        """添加水平线"""
        pass  # 不生成横线

    def save(self, filepath: str) -> None:
        self.document.save(filepath)


def generate_docx(elements: List[Dict[str, Any]],
                  font_settings: FontSettings,
                  paragraph_settings: ParagraphSettings,
                  output_path: str) -> None:
    """从元素列表生成DOCX文档"""
    generator = DocxGenerator(font_settings, paragraph_settings)
    generator.set_document_properties()

    for element in elements:
        elem_type = element.get('type')

        if hasattr(elem_type, 'value'):
            type_str = elem_type.value
        else:
            type_str = str(elem_type)

        if type_str == 'paragraph':
            generator.add_paragraph(element.get('text', ''))
        elif type_str == 'heading':
            generator.add_heading(element.get('text', ''), element.get('level', 1))
        elif type_str == 'code_block':
            generator.add_code_block(element.get('code', ''), element.get('language', ''))
        elif type_str == 'list_item':
            generator.add_list_item(element.get('text', ''))
        elif type_str == 'blockquote':
            generator.add_blockquote(element.get('text', ''))
        elif type_str == 'horizontal_rule':
            generator.add_horizontal_rule()

    generator.save(output_path)


def generate_docx_with_settings(elements: List[Dict[str, Any]],
                                 settings: Dict[str, Any],
                                 output_path: str) -> None:
    """从元素列表生成DOCX文档（支持多种字体和段落设置）"""
    generator = AdvancedDocxGenerator(settings)
    generator.set_document_properties()

    for element in elements:
        elem_type = element.get('type')

        if hasattr(elem_type, 'value'):
            type_str = elem_type.value
        else:
            type_str = str(elem_type)

        if type_str == 'paragraph':
            generator.add_paragraph(element.get('text', ''))
        elif type_str == 'heading':
            generator.add_heading(element.get('text', ''), element.get('level', 1))
        elif type_str == 'code_block':
            generator.add_code_block(element.get('code', ''), element.get('language', ''))
        elif type_str == 'list_item':
            generator.add_list_item(element.get('text', ''))
        elif type_str == 'blockquote':
            generator.add_blockquote(element.get('text', ''))
        elif type_str == 'horizontal_rule':
            generator.add_horizontal_rule()
        elif type_str == 'table':
            generator.add_table(element)
        elif type_str == 'image':
            generator.add_image(element)
        elif type_str == 'formula_block':
            generator.add_formula_block(element.get('formula', ''))

    generator.save(output_path)


class AdvancedDocxGenerator:
    """高级DOCX文档生成器，支持多种字体和段落设置"""

    def __init__(self, settings: Dict[str, Any]):
        self.settings = settings
        self.document = Document()

        self.heading_fonts = settings.get('heading_fonts', [FontSettings() for _ in range(9)])
        self.body_font = settings.get('body_font', FontSettings())
        self.table_font = settings.get('table_font', FontSettings())
        self.caption_font = settings.get('caption_font', FontSettings())

        self.heading_para = settings.get('heading_para', ParagraphSettings())
        self.body_para = settings.get('body_para', ParagraphSettings())
        self.table_para = settings.get('table_para', ParagraphSettings())
        self.caption_para = settings.get('caption_para', ParagraphSettings())
        self.caption_alignment = settings.get('caption_alignment', 'center')

        # 转换选项
        self.remove_spaces = settings.get('remove_spaces', True)
        self.convert_formulas = settings.get('convert_formulas', True)

    def set_document_properties(self):
        """设置文档默认属性"""
        self.document.styles['Normal'].font.name = self.body_font.english_font
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                                 self.body_font.chinese_font)
        self.document.styles['Normal'].font.size = Pt(self.body_font.font_size)
        self.document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        # 初始化标题计数器
        self.heading_counters = [0] * 9

    def apply_font_to_run(self, run, font_setting: FontSettings):
        """应用字体设置到run"""
        run.font.name = font_setting.english_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_setting.chinese_font)
        run.font.size = Pt(font_setting.font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def apply_paragraph_format(self, paragraph, para_setting: ParagraphSettings, is_heading: bool = False):
        """应用段落格式"""
        # 设置行间距
        if para_setting.line_spacing_type == WD_LINE_SPACING.MULTIPLE:
            paragraph.paragraph_format.line_spacing = para_setting.line_spacing
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        else:
            paragraph.paragraph_format.line_spacing_rule = para_setting.line_spacing_type
            paragraph.paragraph_format.line_spacing = para_setting.line_spacing

        if is_heading:
            # 标题：段前段后间距
            paragraph.paragraph_format.space_before = para_setting.heading_space_before
            paragraph.paragraph_format.space_after = para_setting.heading_space_after
        else:
            # 正文：首行缩进，段落之间无额外间距
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # 设置两端对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if para_setting.first_line_indent_enabled:
                # 首行缩进：1 字符 ≈ 字号(pt) 宽度
                indent_value = Pt(para_setting.first_line_indent * self.body_font.font_size)
                paragraph.paragraph_format.first_line_indent = indent_value

    def add_heading(self, text: str, level: int):
        """添加标题（带自动编号）"""
        # 处理引号转换
        text = convert_quotes(text)

        # 去除原有编号（必须在清理空格之前，因为空格是编号的一部分）
        text = self._remove_existing_numbering(text)

        # 处理空格清理
        if self.remove_spaces:
            text = cleanup_text(text)

        if level < 1:
            level = 1
        if level > 9:
            level = 9

        # 更新计数器
        self.heading_counters[level - 1] += 1
        for i in range(level, 9):
            self.heading_counters[i] = 0

        # 生成编号
        number_parts = [str(self.heading_counters[i]) for i in range(level)]
        number_str = '.'.join(number_parts) + ' '

        font_setting = self.heading_fonts[level - 1]
        para_setting = self.heading_para

        p = self.document.add_heading('', level=level if level <= 6 else 6)

        run = p.add_run(number_str + text)
        self.apply_font_to_run(run, font_setting)

        self.apply_paragraph_format(p, para_setting, is_heading=True)

    def _remove_existing_numbering(self, text: str) -> str:
        """去除标题原有的编号"""
        import re
        patterns = [
            # 中文括号编号：（一）、（1）
            r'^[\（\(][一二三四五六七八九十百千零\d]+[\）\)][\s、.]*',
            # 拉丁字母括号编号：（a）、（A）
            r'^[\（\(][a-zA-Zα-ωΑ-Ω]+[\）\)][\s、.]*',
            # 中文数字+标点：一、二、三、
            r'^[一二三四五六七八九十百千零]+[、.][\s]*',
            # 多级数字编号：1.、1.1、1.1.1、1-
            r'^\d+(\.\d+)*[-.、\s]+',
            # 单级数字编号：1、2、3、
            r'^\d+[.、)）\s]+',
            # 拉丁字母编号：a.、A.、a)
            r'^[a-zA-Z][.、)）\s]+',
            # 第几章/第几节：第一章、第一节
            r'^第[一二三四五六七八九十百千\d]+[节章节]?[\s]*',
        ]
        for pattern in patterns:
            text = re.sub(pattern, '', text)
        return text.strip()

    def add_paragraph(self, text: str):
        """添加正文段落"""
        text = convert_quotes(text)

        if self.remove_spaces:
            text = cleanup_text(text)

        if self.convert_formulas:
            parts = self._split_inline_formulas(text)
        else:
            parts = [{'type': 'text', 'text': text}]

        p = self.document.add_paragraph()
        self.apply_paragraph_format(p, self.body_para)

        for part in parts:
            if part['type'] == 'text':
                if part['text'].strip():
                    run = p.add_run(part['text'])
                    self.apply_font_to_run(run, self.body_font)
            elif part['type'] == 'formula':
                self._add_inline_formula(p, part['formula'])

    def _split_inline_formulas(self, text: str) -> List[Dict[str, str]]:
        """分割行内公式和文本"""
        parts = []
        # 只匹配单 $...$ 行内公式，不匹配 $$...$$ 块公式
        pattern = r'(?<!\$)\$([^\$]+?)\$(?!\$)'
        fragments = re.split(pattern, text)

        for i, fragment in enumerate(fragments):
            if i % 2 == 1:
                # 这是公式部分
                formula = fragment
                parts.append({'type': 'formula', 'formula': formula})
            elif fragment.strip():
                parts.append({'type': 'text', 'text': fragment})

        return parts

    def _add_inline_formula(self, paragraph, formula: str):
        """添加行内公式"""
        try:
            mathml_xml = latex_to_mathml(formula.strip())
            mathml_elem = etree.fromstring(mathml_xml)

            # 创建 oMath 元素
            oMath = OxmlElement('m:oMath')
            self._mathml_to_omml(mathml_elem, oMath)

            # 直接将 oMath 添加到段落中，而不是包装在 w:r 中
            paragraph._p.append(oMath)

        except Exception as e:
            run = paragraph.add_run(f'[{formula}]')
            run.font.name = 'Cambria Math'

    def _mathml_to_omml(self, mathml_elem, oMath):
        """将MathML元素转换为OMML"""
        for child in mathml_elem:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag in ('mrow', 'math'):
                self._mathml_to_omml(child, oMath)
            elif tag == 'mi':
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mn':
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mo':
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mtext':
                text = child.text or ''
                r = self._create_omml_run(text)
                oMath.append(r)
            elif tag == 'mspace':
                pass
            elif tag == 'msqrt':
                self._convert_msqrt_simple(child, oMath)
            elif tag == 'mroot':
                self._convert_mroot_simple(child, oMath)
            elif tag == 'mfrac':
                self._convert_mfrac_simple(child, oMath)
            elif tag == 'msub':
                self._convert_msub_simple(child, oMath)
            elif tag == 'msup':
                self._convert_msup_simple(child, oMath)
            elif tag == 'msubsup':
                self._convert_msubsup_simple(child, oMath)
            elif tag == 'munder':
                self._convert_munder_simple(child, oMath)
            elif tag == 'mover':
                self._convert_mover_simple(child, oMath)
            elif tag == 'munderover':
                self._convert_munderover_simple(child, oMath)
            elif tag == 'mrow':
                self._mathml_to_omml(child, oMath)

    def _create_omml_run(self, text: str) -> OxmlElement:
        """创建OMML run元素"""
        r = OxmlElement('m:r')
        rPr = OxmlElement('m:rPr')
        rFonts = OxmlElement('m:rFonts')
        rFonts.set(qn('w:ascii'), 'Cambria Math')
        rFonts.set(qn('w:hAnsi'), 'Cambria Math')
        rFonts.set(qn('w:eastAsia'), 'Cambria Math')
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('m:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        return r

    def _create_ctrlPr(self) -> OxmlElement:
        """创建m:ctrlPr元素，用于公式控制属性"""
        ctrlPr = OxmlElement('m:ctrlPr')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('m:rFonts')
        rFonts.set(qn('w:ascii'), 'Cambria Math')
        rFonts.set(qn('w:hAnsi'), 'Cambria Math')
        rFonts.set(qn('w:eastAsia'), 'Cambria Math')
        rPr.append(rFonts)
        ctrlPr.append(rPr)
        return ctrlPr

    def _convert_msqrt_simple(self, msqrt_elem, oMath):
        """转换平方根"""
        rad = OxmlElement('m:rad')
        deg = OxmlElement('m:deg')
        deg.append(self._create_omml_run('2'))
        rad.append(deg)
        e = OxmlElement('m:e')
        for child in msqrt_elem:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'mrow':
                self._mathml_to_omml(child, e)
        rad.append(e)
        oMath.append(rad)

    def _convert_mroot_simple(self, mroot_elem, oMath):
        """转换根号"""
        rad = OxmlElement('m:rad')
        children = list(mroot_elem)
        if len(children) >= 2:
            deg = OxmlElement('m:deg')
            self._mathml_to_omml(children[1], deg)
            rad.append(deg)
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            rad.append(e)
            oMath.append(rad)

    def _convert_mfrac_simple(self, mfrac_elem, oMath):
        """转换分数"""
        f = OxmlElement('m:f')
        children = list(mfrac_elem)
        if len(children) >= 2:
            num = OxmlElement('m:num')
            self._mathml_to_omml(children[0], num)
            f.append(num)
            den = OxmlElement('m:den')
            self._mathml_to_omml(children[1], den)
            f.append(den)
            e = OxmlElement('m:e')
            e.append(f)
            oMath.append(e)

    def _convert_msub_simple(self, msub_elem, oMath):
        """转换下标"""
        children = list(msub_elem)
        if len(children) >= 2:
            sSub = OxmlElement('m:sSub')
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            sSub.append(e)
            sub = OxmlElement('m:sub')
            self._mathml_to_omml(children[1], sub)
            sSub.append(sub)
            oMath.append(sSub)

    def _convert_msup_simple(self, msup_elem, oMath):
        """转换上标"""
        children = list(msup_elem)
        if len(children) >= 2:
            sSup = OxmlElement('m:sSup')
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            sSup.append(e)
            sup = OxmlElement('m:sup')
            self._mathml_to_omml(children[1], sup)
            sSup.append(sup)
            oMath.append(sSup)

    def _convert_msubsup_simple(self, msubsup_elem, oMath):
        """转换上下标"""
        children = list(msubsup_elem)
        if len(children) >= 3:
            sSubSup = OxmlElement('m:sSubSup')
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            sSubSup.append(e)
            sub = OxmlElement('m:sub')
            self._mathml_to_omml(children[1], sub)
            sSubSup.append(sub)
            sup = OxmlElement('m:sup')
            self._mathml_to_omml(children[2], sup)
            sSubSup.append(sup)
            oMath.append(sSubSup)

    def _convert_munder_simple(self, munder_elem, oMath):
        """转换下划线"""
        children = list(munder_elem)
        if len(children) >= 2:
            acc = OxmlElement('m:acc')
            accPr = OxmlElement('m:accPr')
            chr_elem = OxmlElement('m:chr')
            chr_elem.set(qn('m:val'), '_')
            accPr.append(chr_elem)
            ctrlPr = self._create_ctrlPr()
            accPr.append(ctrlPr)
            acc.append(accPr)
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            acc.append(e)
            oMath.append(acc)

    def _convert_mover_simple(self, mover_elem, oMath):
        """转换上划线"""
        children = list(mover_elem)
        if len(children) >= 2:
            acc = OxmlElement('m:acc')
            accPr = OxmlElement('m:accPr')
            chr_elem = OxmlElement('m:chr')
            chr_elem.set(qn('m:val'), '‾')
            accPr.append(chr_elem)
            ctrlPr = self._create_ctrlPr()
            accPr.append(ctrlPr)
            acc.append(accPr)
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            acc.append(e)
            oMath.append(acc)

    def _convert_munderover_simple(self, munderover_elem, oMath):
        """转换上下组合"""
        children = list(munderover_elem)
        if len(children) >= 3:
            acc = OxmlElement('m:acc')
            accPr = OxmlElement('m:accPr')
            chr_elem = OxmlElement('m:chr')
            chr_elem.set(qn('m:val'), '‾')
            accPr.append(chr_elem)
            ctrlPr = self._create_ctrlPr()
            accPr.append(ctrlPr)
            acc.append(accPr)
            e = OxmlElement('m:e')
            self._mathml_to_omml(children[0], e)
            acc.append(e)
            sub = OxmlElement('m:sub')
            self._mathml_to_omml(children[1], sub)
            acc.append(sub)
            sup = OxmlElement('m:sup')
            self._mathml_to_omml(children[2], sup)
            acc.append(sup)
            oMath.append(acc)

    def add_list_item(self, text: str):
        """添加列表项"""
        text = cleanup_text(text)
        text = convert_quotes(text)

        p = self.document.add_paragraph()
        self.apply_paragraph_format(p, self.body_para)

        run = p.add_run(f"• {text}")
        self.apply_font_to_run(run, self.body_font)

    def add_formula_block(self, formula: str):
        """添加块公式：转换为 docx 可编辑的公式，居中显示"""
        formula = (formula or '').strip()
        if not formula:
            return

        # 使用 docx 的 add_paragraph 创建段落，并设置居中对齐
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 块公式不应用正文的首行缩进
        p.paragraph_format.first_line_indent = Pt(0)
        # 块公式前后留出适当空间
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 段间行间距使用正文的设置
        if self.body_para.line_spacing_type == WD_LINE_SPACING.MULTIPLE:
            p.paragraph_format.line_spacing = self.body_para.line_spacing
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        try:
            # latex -> mathml -> OMML
            mathml_xml = latex_to_mathml(formula)
            mathml_elem = etree.fromstring(mathml_xml)
            oMath = OxmlElement('m:oMath')
            self._mathml_to_omml(mathml_elem, oMath)

            # oMathPara：让 Word 将其识别为独立的可编辑公式段落（双击可进入公式编辑器）
            oMathPara = OxmlElement('m:oMathPara')
            oMathPara.append(oMath)

            # 段落内需要至少一个 w:r 来承载 oMathPara；直接添加到 p 元素
            p._p.append(oMathPara)
        except Exception:
            # 转换失败：降级为普通文本
            run = p.add_run(formula)
            self.apply_font_to_run(run, self.body_font)

    def add_code_block(self, code: str, language: str):
        """添加代码块"""
        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)

        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._p.get_or_add_pPr().append(shading)

    def add_blockquote(self, text: str):
        """添加引用块"""
        text = cleanup_text(text)
        text = convert_quotes(text)

        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)

        run = p.add_run(text)
        run.font.italic = True
        self.apply_font_to_run(run, self.body_font)

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '808080')
        pBdr.append(left)
        pPr.append(pBdr)

    def add_horizontal_rule(self):
        """添加水平线（默认跳过，不生成横线）"""
        pass

    def add_table(self, element: Dict[str, Any]):
        """添加表格"""
        headers = element.get('headers', [])
        rows = element.get('rows', [])

        if not headers and not rows:
            return

        table = self.document.add_table(rows=0, cols=len(headers) if headers else len(rows[0]) if rows else 0)
        table.style = 'Table Grid'

        if headers:
            header_row = table.add_row()
            for i, header_text in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = cleanup_text(header_text)
                self.apply_font_to_run(cell.paragraphs[0].runs[0], self.table_font)
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                cell = row.cells[i]
                cell.text = cleanup_text(cell_text)
                if i < len(cell.paragraphs[0].runs):
                    self.apply_font_to_run(cell.paragraphs[0].runs[0], self.table_font)

        self.apply_paragraph_format(table.rows[0].cells[0].paragraphs[0], self.table_para)

    def add_image(self, element: Dict[str, Any]):
        """添加图片（图名）"""
        caption_text = element.get('caption', '')
        if caption_text:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if self.caption_alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(caption_text)
            self.apply_font_to_run(run, self.caption_font)
            self.apply_paragraph_format(p, self.caption_para)

    def save(self, filepath: str):
        """保存文档"""
        self.document.save(filepath)
